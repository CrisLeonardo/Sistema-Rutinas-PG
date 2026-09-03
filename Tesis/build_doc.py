# -*- coding: utf-8 -*-
"""Construye el informe de Proyecto de Graduacion I - Cuarta Entrega (APA, TNR 12, 1.5)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGS = "_figs"
OUTNAME = "Leonardo PG1 Entrega 4.docx"

doc = Document()

# ----------------------------------------------------------------------------
# CONFIGURACION GLOBAL: fuente, interlineado, margenes, estilos APA
# ----------------------------------------------------------------------------
def set_font(style, name="Times New Roman", size=12, bold=None, italic=None,
             color=None, caps=False):
    f = style.font
    f.name = name
    f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), name)
    if caps:
        c = OxmlElement('w:caps'); c.set(qn('w:val'), 'true'); rpr.append(c)

# Estilo Normal
normal = doc.styles['Normal']
set_font(normal, size=12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(6)
pf.space_before = Pt(0)

# Estilos de titulo APA
h1 = doc.styles['Heading 1']
set_font(h1, size=14, bold=True, color=RGBColor(0, 0, 0), caps=True)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(12)
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
h1.paragraph_format.keep_with_next = True

h2 = doc.styles['Heading 2']
set_font(h2, size=12, bold=True, color=RGBColor(0, 0, 0))
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(10)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
h2.paragraph_format.keep_with_next = True

h3 = doc.styles['Heading 3']
set_font(h3, size=12, bold=True, italic=True, color=RGBColor(0, 0, 0))
h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(4)
h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
h3.paragraph_format.keep_with_next = True

# Margenes 3-3-3-2 cm
sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.bottom_margin = Cm(3)
sec.left_margin = Cm(3)
sec.right_margin = Cm(2)

# Numero de pagina en el pie (centrado); ocultar en caratula (primera pagina distinta)
sec.different_first_page_header_footer = True

def add_page_number(paragraph):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = "PAGE"
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "1"
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for x in (b, i, s, t, e):
        run._r.append(x)
    run.font.name = "Times New Roman"; run.font.size = Pt(12)

footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(fp)

# ----------------------------------------------------------------------------
# HELPERS DE CONTENIDO
# ----------------------------------------------------------------------------
INDENT = Cm(1.25)

def H1(text):
    return doc.add_heading(text, level=1)

def H2(text):
    return doc.add_heading(text, level=2)

def H3(text):
    return doc.add_heading(text, level=3)

def P(text, indent=True, justify=True, after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = INDENT
    return p

def DASH(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("− " + text)
    return p

def LETTER(letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(letter + ") ")
    r.bold = True
    p.add_run(text)
    return p

def page_break():
    doc.add_page_break()

def field(paragraph, instr, default="Actualice este campo en Word (clic derecho > Actualizar campos)."):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = default
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for x in (b, i, s, t, e):
        run._r.append(x)
    return run

def seq(paragraph, label):
    """Inserta un campo SEQ (numeracion automatica) para Tabla/Figura."""
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve')
    i.text = ' SEQ %s \\* ARABIC ' % label
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "1"
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for x in (b, i, s, t, e):
        run._r.append(x)

def fig_caption(title):
    """Pie de figura: 'Figura N. titulo' (con SEQ para indice de figuras)."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style = doc.styles['Caption']
    r = p.add_run("Figura ")
    r.bold = True
    seq(p, "Figura")
    rr = p.add_run(". " + title)
    rr.bold = True
    return p

def tbl_caption(title):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style = doc.styles['Caption']
    r = p.add_run("Tabla ")
    r.bold = True
    seq(p, "Tabla")
    rr = p.add_run(". " + title)
    rr.bold = True
    return p

def source(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("Nota. ")
    r.italic = True; r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.italic = True; r2.font.size = Pt(10)
    return p

def figure(filename, title, src, width_cm=12.5):
    from docx.shared import Cm as _Cm
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    path = os.path.join(FIGS, filename)
    run.add_picture(path, width=_Cm(width_cm))
    fig_caption(title)
    source(src)

# Ensure Caption style exists
if 'Caption' not in [s.name for s in doc.styles]:
    pass  # docx tiene 'Caption' por defecto
cap = doc.styles['Caption']
set_font(cap, size=11, italic=False, color=RGBColor(0, 0, 0))

print("infraestructura lista")

# ----------------------------------------------------------------------------
# CONTENIDO
# ----------------------------------------------------------------------------
import content
content.build(globals())

# ----------------------------------------------------------------------------
# Actualizar campos al abrir
# ----------------------------------------------------------------------------
settings = doc.settings.element
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
settings.append(uf)

doc.save(OUTNAME)
print("GUARDADO:", OUTNAME)
